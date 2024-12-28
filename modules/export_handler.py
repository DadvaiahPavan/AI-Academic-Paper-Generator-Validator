import docx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os
import io

def export_document(draft_content, research_topic):
    """
    Export the academic draft to DOCX and PDF formats
    
    Args:
        draft_content (str): Content of the academic draft
        research_topic (str): Research topic for file naming
    
    Returns:
        dict: Paths to exported files
    """
    # Ensure export directory exists
    export_dir = os.path.join(os.getcwd(), 'exports')
    os.makedirs(export_dir, exist_ok=True)

    # Sanitize filename
    safe_filename = "".join(
        c for c in research_topic 
        if c.isalnum() or c in (' ', '_')
    ).rstrip()

    # Export to DOCX
    docx_path = os.path.join(export_dir, f"{safe_filename}_draft.docx")
    doc = docx.Document()
    
    # Handle potential encoding issues
    try:
        # Ensure draft_content is a string and clean any problematic characters
        draft_content = str(draft_content).encode('utf-8', errors='replace').decode('utf-8')
        doc.add_paragraph(draft_content)
    except Exception as e:
        print(f"Error adding paragraph to DOCX: {e}")
        doc.add_paragraph("Error generating draft content")
    
    doc.save(docx_path)

    # Export to PDF
    pdf_path = os.path.join(export_dir, f"{safe_filename}_draft.pdf")
    
    # Register a font to handle potential encoding issues
    pdfmetrics.registerFont(TTFont('Arial', 'Arial.ttf'))
    
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    
    # Set font and prepare text
    c.setFont('Arial', 12)
    text_object = c.beginText(50, height - 50)
    
    # Handle potential encoding issues when writing to PDF
    try:
        # Split content into lines, handling potential encoding problems
        lines = draft_content.split('\n')
        for line in lines:
            # Encode and decode to handle potential problematic characters
            clean_line = line.encode('utf-8', errors='replace').decode('utf-8')
            text_object.textLine(clean_line)
    except Exception as e:
        print(f"Error writing PDF content: {e}")
        text_object.textLine("Error generating draft content")
    
    c.drawText(text_object)
    c.showPage()
    c.save()

    return {
        'docx_path': docx_path,
        'pdf_path': pdf_path
    }
